from typing import List, Dict, Optional
import logging
import os
import hashlib
import re
from collections import Counter

from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document as DocxDocument
import markdown


from app.config import settings

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentStore:
    def __init__(self):
        self.documents: List[Document] = []
        self.embeddings = OllamaEmbeddings(model="llama3")
        
        # ChromaDB persistent storage
        self.chroma_dir = "chroma_db"
        self.vector_store = Chroma(
            persist_directory=self.chroma_dir,
            embedding_function=self.embeddings,
            collection_name="documents"
        )
        
        # 시작 시 기존 문서 개수 확인
        self._load_existing_documents()
        
    def _load_existing_documents(self):
        """ChromaDB에서 기존 문서 개수 확인"""
        try:
            collection = self.vector_store._collection
            count = collection.count()
            logger.info(f"ChromaDB에서 {count}개의 기존 문서 발견")
        except Exception as e:
            logger.info(f"ChromaDB 초기화: {e}")
    
    def _document_exists(self, filename: str, file_hash: str) -> bool:
        """문서가 이미 ChromaDB에 있는지 확인"""
        try:
            # 메타데이터로 검색 (AND 조건)
            results = self.vector_store.get(
                where={
                    "$and": [
                        {"filename": {"$eq": filename}},
                        {"file_hash": {"$eq": file_hash}}
                    ]
                }
            )
            return len(results['ids']) > 0
        except Exception as e:
            logger.error(f"문서 존재 확인 실패: {e}")
            return False
    
    def _get_file_hash(self, file_path: str) -> str:
        """파일의 해시값을 계산"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    
    def _load_pdf(self, file_path: str, filename: str, file_hash: str) -> List[Document]:
        """
        PDF 파일을 페이지별로 로딩하여 Document 객체 리스트로 반환
        
        Args:
            file_path (str): PDF 파일의 경로
            filename (str): 원본 파일명
            file_hash (str): 파일 해시값 (중복 방지용)
            
        Returns:
            List[Document]: 각 페이지를 Document 객체로 변환한 리스트
            
        Features:
            - PyPDFLoader를 사용하여 PDF를 페이지별로 분할
            - 각 페이지에 상세한 메타데이터 추가 (파일명, 해시, 페이지 번호 등)
            - 출처 정보에 페이지 번호 포함으로 정확한 위치 추적 가능
        """
        loader = PyPDFLoader(file_path)
        pages = []
        
        logger.info("PDF 페이지 로딩 중...")
        for i, page in enumerate(loader.load()):
            # 각 페이지에 상세한 메타데이터 추가
            page.metadata.update({
                'filename': filename,           # 원본 파일명
                'file_hash': file_hash,        # 파일 해시 (중복 방지)
                'page_number': i + 1,          # 페이지 번호 (1부터 시작)
                'file_type': 'pdf',            # 파일 타입
                'source': f"{filename} (페이지 {i + 1})"  # 사용자에게 보여질 출처 정보
            })
            pages.append(page)
            logger.info(f"페이지 {i+1} 로딩 완료")
        
        return pages
    
    def _load_docx(self, file_path: str, filename: str, file_hash: str) -> List[Document]:
        """
        DOCX 파일을 청크 단위로 로딩하여 Document 객체 리스트로 반환
        
        Args:
            file_path (str): DOCX 파일의 경로
            filename (str): 원본 파일명  
            file_hash (str): 파일 해시값 (중복 방지용)
            
        Returns:
            List[Document]: 청크별로 분할된 Document 객체 리스트
            
        Features:
            - python-docx를 사용하여 Word 문서의 단락 추출
            - 빈 단락 자동 제거로 불필요한 공백 방지
            - 500자 기준으로 적절한 크기의 청크 생성
            - 각 청크에 번호와 위치 정보 메타데이터 추가
        """
        logger.info("DOCX 문서 로딩 중...")
        doc = DocxDocument(file_path)
        
        # 단락별로 분할하고 빈 단락 제거
        paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 빈 단락 제외
                paragraphs.append(paragraph.text.strip())
        
        # 여러 단락을 묶어서 적절한 크기의 청크 생성 (대략 500자씩)
        chunks = []
        current_chunk = ""
        chunk_num = 1
        
        for paragraph in paragraphs:
            # 현재 청크에 단락을 추가했을 때 500자를 초과하면 새 청크 생성
            if len(current_chunk + paragraph) > 500 and current_chunk:
                chunks.append(Document(
                    page_content=current_chunk.strip(),
                    metadata={
                        'filename': filename,              # 원본 파일명
                        'file_hash': file_hash,           # 파일 해시
                        'chunk_number': chunk_num,        # 청크 번호
                        'file_type': 'docx',             # 파일 타입
                        'source': f"{filename} (청크 {chunk_num})"  # 출처 정보
                    }
                ))
                current_chunk = paragraph + "\n"
                chunk_num += 1
            else:
                current_chunk += paragraph + "\n"
        
        # 마지막 청크 추가 (남은 내용이 있는 경우)
        if current_chunk.strip():
            chunks.append(Document(
                page_content=current_chunk.strip(),
                metadata={
                    'filename': filename,
                    'file_hash': file_hash,
                    'chunk_number': chunk_num,
                    'file_type': 'docx',
                    'source': f"{filename} (청크 {chunk_num})"
                }
            ))
        
        logger.info(f"총 {len(chunks)}개 청크 생성")
        return chunks
    
    def _load_text(self, file_path: str, filename: str, file_hash: str) -> List[Document]:
        """
        TXT/MD 파일을 청크 단위로 로딩하여 Document 객체 리스트로 반환
        
        Args:
            file_path (str): 텍스트 파일의 경로
            filename (str): 원본 파일명
            file_hash (str): 파일 해시값 (중복 방지용)
            
        Returns:
            List[Document]: 청크별로 분할된 Document 객체 리스트
            
        Features:
            - UTF-8 인코딩으로 텍스트 파일 읽기
            - 마크다운 파일의 경우 HTML 변환 후 태그 제거
            - 1000자 단위로 청크 분할 (검색 효율성과 정확도 균형)
            - 각 청크에 순서 정보와 파일 타입 메타데이터 추가
        """
        logger.info(f"텍스트 파일 로딩 중: {filename}")
        
        # UTF-8 인코딩으로 파일 내용 읽기
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 마크다운 파일 처리: HTML로 변환 후 태그 제거하여 순수 텍스트 추출
        if filename.endswith('.md'):
            html = markdown.markdown(content)
            # 간단한 정규식으로 HTML 태그 제거
            import re
            content = re.sub('<[^<]+?>', '', html)
        
        # 긴 텍스트를 청크로 분할 (1000자씩)
        # 1000자는 검색 정확도와 컨텍스트 유지의 적절한 균형점
        chunks = []
        chunk_size = 1000
        
        for i in range(0, len(content), chunk_size):
            chunk_content = content[i:i + chunk_size]
            if chunk_content.strip():  # 빈 청크 제외
                chunks.append(Document(
                    page_content=chunk_content.strip(),
                    metadata={
                        'filename': filename,                                    # 원본 파일명
                        'file_hash': file_hash,                                 # 파일 해시
                        'chunk_number': i // chunk_size + 1,                   # 청크 순서 번호
                        'file_type': 'txt' if filename.endswith('.txt') else 'md',  # 파일 타입 구분
                        'source': f"{filename} (청크 {i // chunk_size + 1})"   # 출처 정보
                    }
                ))
        
        logger.info(f"총 {len(chunks)}개 청크 생성")
        return chunks

    def add_document(self, file_path: str, original_filename: str = None) -> None:
        """문서를 저장소에 추가"""
        filename = original_filename or os.path.basename(file_path)
        logger.info(f"문서 로딩 시작: {filename}")
        
        # 파일 해시 계산
        file_hash = self._get_file_hash(file_path)
        
        # ChromaDB에 이미 존재하는지 확인
        if self._document_exists(filename, file_hash):
            logger.info(f"문서가 이미 ChromaDB에 존재함: {filename}")
            return
        
        # 파일 형식에 따른 처리
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            if file_ext == '.pdf':
                pages = self._load_pdf(file_path, filename, file_hash)
            elif file_ext == '.docx':
                pages = self._load_docx(file_path, filename, file_hash)
            elif file_ext in ['.txt', '.md']:
                pages = self._load_text(file_path, filename, file_hash)
            else:
                logger.error(f"지원되지 않는 파일 형식: {file_ext}")
                return
                
            logger.info("ChromaDB에 문서 추가 중...")
            self.vector_store.add_documents(documents=pages)
            self.documents.extend(pages)
            logger.info("문서 추가 완료")
            
        except Exception as e:
            logger.error(f"문서 처리 실패: {e}")
            raise
    
    def _keyword_search(self, query: str, documents: List[Document], k: int = 10) -> List[Dict]:
        """
        키워드 기반 검색 (BM25 스타일 스코어링)
        
        Args:
            query (str): 검색 쿼리
            documents (List[Document]): 검색 대상 문서 리스트
            k (int): 반환할 최대 결과 수
            
        Returns:
            List[Dict]: 스코어와 함께 정렬된 문서 리스트
            
        Algorithm:
            - 쿼리를 단어 단위로 분할하여 키워드 추출
            - 각 문서에서 키워드 빈도(TF) 계산
            - TF-IDF 스타일로 관련성 점수 산출
            - 벡터 검색과 상호 보완하여 정확도 향상
        """
        # 쿼리에서 키워드 추출 (알파벳과 숫자만)
        query_terms = re.findall(r'\w+', query.lower())
        if not query_terms:
            return []
        
        scored_docs = []
        for i, doc in enumerate(documents):
            content = doc.page_content.lower()
            content_terms = re.findall(r'\w+', content)
            
            if not content_terms:
                continue
            
            # TF-IDF 스타일 스코어 계산
            doc_term_freq = Counter(content_terms)  # 문서 내 단어 빈도 계산
            score = 0
            
            for term in query_terms:
                if term in doc_term_freq:
                    # TF (Term Frequency): 단어 빈도 / 전체 단어 수
                    tf = doc_term_freq[term] / len(content_terms)
                    # IDF는 간단히 1.0으로 설정 (실제로는 전체 문서 컬렉션 기반 계산)
                    idf = 1.0
                    score += tf * idf
            
            # 점수가 0보다 큰 문서만 결과에 포함
            if score > 0:
                scored_docs.append({
                    'document': doc,
                    'score': score,
                    'index': i
                })
        
        # 점수 내림차순으로 정렬하고 상위 k개 반환
        scored_docs.sort(key=lambda x: x['score'], reverse=True)
        return scored_docs[:k]
    
    def search_similar(self, query: str, k: Optional[int] = 1):
        """
        하이브리드 검색: 벡터 검색 + 키워드 검색 + 리랭킹 시스템
        
        Args:
            query (str): 검색 쿼리
            k (Optional[int]): 반환할 결과 수 (기본값: 1)
            
        Returns:
            List[Document]: 관련성이 높은 순서로 정렬된 문서 리스트
            
        Algorithm:
            1. 벡터 검색: 의미적 유사성 기반 검색 (임베딩 벡터 유사도)
            2. 키워드 검색: 정확한 용어 매칭 기반 검색 (TF-IDF)
            3. 하이브리드 스코어링: 벡터(70%) + 키워드(30%) 가중 평균
            4. 리랭킹: 구문 일치, 근접성, 문서 품질 등을 고려한 재순위 매기기
            
        Benefits:
            - 의미적 검색과 정확한 용어 매칭을 동시에 활용
            - 단일 방법의 한계를 상호 보완
            - 리랭킹으로 최종 결과의 관련성 최적화
        """
        logger.info(f"'{query}' 하이브리드 검색 중... (k={k})")
        
        # 1. 벡터 검색: 의미적 유사성 기반 (더 많은 후보 확보를 위해 k*2개)
        vector_docs = self.vector_store.similarity_search_with_score(query=query, k=k*2)
        logger.info(f"벡터 검색: {len(vector_docs)}개 문서 발견")
        
        # 2. 키워드 검색: 정확한 용어 매칭 기반
        try:
            # ChromaDB에서 모든 문서 데이터 가져오기
            all_docs_data = self.vector_store.get()
            all_documents = []
            for metadata, content in zip(all_docs_data['metadatas'], all_docs_data['documents']):
                doc = Document(page_content=content, metadata=metadata)
                all_documents.append(doc)
            
            keyword_results = self._keyword_search(query, all_documents, k*2)
            logger.info(f"키워드 검색: {len(keyword_results)}개 문서 발견")
            
        except Exception as e:
            # 키워드 검색 실패 시 벡터 검색 결과만 반환
            logger.warning(f"키워드 검색 실패, 벡터 검색만 사용: {e}")
            return [doc for doc, _ in vector_docs]
        
        # 3. 하이브리드 스코어 계산: 벡터 검색과 키워드 검색 결과 통합
        hybrid_scores = {}
        
        # 3-1. 벡터 검색 점수 처리
        for doc, sim_score in vector_docs:
            # 문서 고유 식별자 생성 (파일명 + 내용 일부)
            doc_key = f"{doc.metadata.get('filename', 'unknown')}_{doc.page_content[:100]}"
            # 유사도 거리를 점수로 변환 (거리가 작을수록 높은 점수)
            vector_score = 1.0 / (1.0 + sim_score) if sim_score > 0 else 1.0
            hybrid_scores[doc_key] = {
                'document': doc,
                'vector_score': vector_score,
                'keyword_score': 0.0
            }
        
        # 3-2. 키워드 검색 점수 추가
        for result in keyword_results:
            doc = result['document']
            doc_key = f"{doc.metadata.get('filename', 'unknown')}_{doc.page_content[:100]}"
            
            if doc_key in hybrid_scores:
                # 벡터 검색에서 발견된 문서의 키워드 점수 추가
                hybrid_scores[doc_key]['keyword_score'] = result['score']
            else:
                # 키워드 검색에서만 발견된 새로운 문서 추가
                hybrid_scores[doc_key] = {
                    'document': doc,
                    'vector_score': 0.0,
                    'keyword_score': result['score']
                }
        
        # 3-3. 하이브리드 점수 계산 (가중 평균)
        # 벡터 검색 70%, 키워드 검색 30%로 가중치 설정
        vector_weight = 0.7    # 의미적 유사성 중시
        keyword_weight = 0.3   # 정확한 용어 매칭 보조
        
        final_results = []
        for doc_key, scores in hybrid_scores.items():
            # 가중 평균으로 최종 점수 계산
            final_score = (vector_weight * scores['vector_score'] + 
                          keyword_weight * scores['keyword_score'])
            final_results.append({
                'document': scores['document'],
                'score': final_score
            })
        
        # 하이브리드 점수로 1차 정렬
        final_results.sort(key=lambda x: x['score'], reverse=True)
        
        # 4. 리랭킹 적용: 하이브리드 점수를 바탕으로 추가 요소들을 고려한 재순위 매기기
        reranked_results = self._rerank_results(query, final_results[:k*2])
        
        # 5. 최종 결과 반환: 상위 k개 문서 추출
        top_docs = [result['document'] for result in reranked_results[:k]]
        logger.info(f"하이브리드 검색 + 리랭킹 완료: {len(top_docs)}개 최종 결과")
        
        return top_docs
    
    def _rerank_results(self, query: str, results: List[Dict]) -> List[Dict]:
        """
        고급 리랭킹 시스템: 다양한 관련성 지표를 활용한 검색 결과 재순위 매기기
        
        Args:
            query (str): 원본 검색 쿼리
            results (List[Dict]): 하이브리드 검색 결과 리스트
            
        Returns:
            List[Dict]: 리랭킹된 결과 리스트
            
        Reranking Factors:
            1. 정확한 구문 일치 (Exact Phrase Match)
            2. 쿼리 단어의 근접성 (Query Term Proximity) 
            3. 문서 시작 부분 일치 (Document Start Match)
            4. 문서 길이 정규화 (Document Length Normalization)
            5. 파일 타입별 가중치 (File Type Weighting)
            
        Benefits:
            - 하이브리드 점수만으로는 포착하기 어려운 관련성 신호 활용
            - 사용자 의도와 더 정확하게 일치하는 결과 우선 순위화
            - 문서 품질과 구조적 특성을 고려한 개선된 순위
        """
        query_terms = set(re.findall(r'\w+', query.lower()))
        
        for result in results:
            doc = result['document']
            content = doc.page_content.lower()
            
            # 기본 하이브리드 점수에서 시작
            rerank_score = result['score']
            
            # 1. 정확한 구문 일치 보너스 (50% 보너스)
            # 쿼리가 문서에 완전히 포함된 경우 높은 관련성으로 판단
            if query.lower() in content:
                rerank_score *= 1.5
                logger.debug(f"정확한 구문 일치 보너스 적용: {doc.metadata.get('filename', 'unknown')}")
            
            # 2. 쿼리 단어의 근접성 점수
            # 쿼리의 단어들이 문서에서 가까이 나타날수록 높은 점수
            content_words = re.findall(r'\w+', content)
            query_word_positions = []
            
            # 쿼리 단어들의 문서 내 위치 찾기
            for i, word in enumerate(content_words):
                if word in query_terms:
                    query_word_positions.append(i)
            
            if len(query_word_positions) > 1:
                # 단어 간 거리가 가까울수록 높은 근접성 점수
                proximity_score = 0
                for i in range(len(query_word_positions) - 1):
                    distance = query_word_positions[i + 1] - query_word_positions[i]
                    proximity_score += 1.0 / (1.0 + distance)  # 거리 역수로 점수 계산
                
                # 근접성 점수를 리랭킹 점수에 반영 (최대 10% 보너스)
                rerank_score *= (1.0 + proximity_score * 0.1)
            
            # 3. 문서 시작 부분 일치 보너스
            # 문서 시작 부분에 쿼리 단어가 나타나면 중요도가 높다고 판단
            content_start = content[:200]  # 첫 200자 검사
            start_matches = sum(1 for term in query_terms if term in content_start)
            if start_matches > 0:
                # 시작 부분 일치 비율에 따른 보너스 (최대 20% 보너스)
                start_bonus = start_matches / len(query_terms) * 0.2
                rerank_score *= (1.0 + start_bonus)
            
            # 4. 문서 길이 정규화
            # 너무 짧거나 긴 문서는 품질이 낮을 수 있다고 가정
            content_length = len(content.split())
            if content_length < 10:  # 너무 짧은 문서 (10단어 미만)
                rerank_score *= 0.8  # 20% 페널티
            elif content_length > 1000:  # 너무 긴 문서 (1000단어 초과)
                rerank_score *= 0.9  # 10% 페널티
            
            # 5. 파일 타입 기반 가중치 조정
            # 특정 파일 타입에 대한 신뢰도 반영
            file_type = doc.metadata.get('file_type', '').lower()
            if file_type == 'pdf':
                rerank_score *= 1.1   # PDF: 11% 보너스 (공식 문서 가능성)
            elif file_type == 'docx':
                rerank_score *= 1.05  # DOCX: 5% 보너스 (구조화된 문서)
            # TXT, MD 파일은 가중치 변경 없음
            
            # 최종 리랭킹 점수 저장
            result['rerank_score'] = rerank_score
        
        # 리랭킹 점수 기준으로 재정렬 (내림차순)
        results.sort(key=lambda x: x['rerank_score'], reverse=True)
        logger.info(f"리랭킹 완료: {len(results)}개 결과")
        
        return results
    
    def get_all_documents(self) -> List[Document]:
        """모든 문서 반환"""
        return self.documents
    
    def get_document_count(self) -> int:
        """문서 개수 반환"""
        try:
            return self.vector_store._collection.count()
        except:
            return len(self.documents)
    
    def clear_documents(self) -> None:
        """모든 문서 삭제"""
        self.documents.clear()
        try:
            self.vector_store._collection.delete()
            logger.info("ChromaDB에서 모든 문서 삭제됨")
        except Exception as e:
            logger.error(f"ChromaDB 삭제 실패: {e}")

# 전역 문서 저장소 인스턴스
document_store = DocumentStore()